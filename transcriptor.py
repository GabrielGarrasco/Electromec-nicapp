import streamlit as st
import os
import io
import re
from PIL import Image
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Inches

def add_markdown_runs(paragraph, text):
    """Traduce negritas y cursivas de Markdown al formato nativo de Word"""
    # Separamos por negritas (**texto**)
    partes = re.split(r'(\*\*.*?\*\*)', text)
    for part in partes:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Separamos por cursivas (*texto* o _texto_)
            subpartes = re.split(r'(\*[^\*]+\*|_[^_]+_)', part)
            for subpart in subpartes:
                if (subpart.startswith('*') and subpart.endswith('*')) or \
                   (subpart.startswith('_') and subpart.endswith('_')):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(subpart)

def generar_docx(texto, imagenes):
    doc = Document()
    doc.add_heading("Apuntes Digitalizados", 0)
    
    # Procesamos línea por línea asignando estilos nativos de Word
    for linea in texto.split('\n'):
        linea = linea.strip()
        if not linea:
            continue
            
        if linea.startswith('# '):
            doc.add_heading(linea[2:].replace('**', ''), 1)
        elif linea.startswith('## '):
            doc.add_heading(linea[3:].replace('**', ''), 2)
        elif linea.startswith('### '):
            doc.add_heading(linea[4:].replace('**', ''), 3)
        elif linea.startswith('* ') or linea.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_markdown_runs(p, linea[2:])
        elif re.match(r'^\d+\.\s', linea):
            p = doc.add_paragraph(style='List Number')
            texto_lista = re.sub(r'^\d+\.\s', '', linea)
            add_markdown_runs(p, texto_lista)
        else:
            p = doc.add_paragraph()
            add_markdown_runs(p, linea)
    
    # Añadimos las imágenes ordenadas al final del documento
    if imagenes:
        doc.add_page_break()
        doc.add_heading("Imágenes Originales", 1)
        for img_dict in imagenes:
            img = img_dict['img']
            img_byte_arr = io.BytesIO()
            # Convertimos a RGB por si hay PNGs con transparencia
            if img.mode in ("RGBA", "P"): 
                img = img.convert("RGB")
            img.save(img_byte_arr, format='JPEG')
            img_byte_arr.seek(0)
            doc.add_picture(img_byte_arr, width=Inches(6.0))
            doc.add_paragraph(f"Página {img_dict['orden']}")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def renderizar_transcriptor():
    st.header("Digitalizar Apuntes")

    # Conectar con la API Key guardada en los Secrets
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        client = genai.Client(api_key=api_key)
    except Exception as e:
        st.error("⚠️ No se encontró la API Key. Revisá los Secrets en Streamlit Cloud.")
        return
        
    if 'dicc_abreviaturas' not in st.session_state:
        st.session_state['dicc_abreviaturas'] = "Ej: q = que, cto = circuito, tmb = también"

    # Diccionario manual opcional
    with st.expander("Mis Abreviaturas (Opcional)"):
        st.info("Anotá acá tus abreviaturas si querés forzar a la IA a que las lea de una manera específica.")
        st.session_state['dicc_abreviaturas'] = st.text_area("Diccionario", st.session_state['dicc_abreviaturas'], height=100, label_visibility="collapsed")

    # Subir archivos
    archivos_subidos = st.file_uploader("Subí las fotos de tus apuntes acá", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

    if archivos_subidos:
        st.subheader("Ordená tus páginas")
        
        cols = st.columns(len(archivos_subidos))
        imagenes_ordenadas = []
        
        for i, archivo in enumerate(archivos_subidos):
            img = Image.open(archivo)
            with cols[i]:
                st.image(img, use_container_width=True)
                orden = st.number_input("Página", min_value=1, value=i+1, key=f"orden_{archivo.name}")
                imagenes_ordenadas.append({"orden": orden, "img": img})

        imagenes_ordenadas.sort(key=lambda x: x["orden"])
        imagenes_pil_final = [item["img"] for item in imagenes_ordenadas]

        st.divider()

        # Botón único de proceso
        if st.button("🧠 Procesar y Transcribir Todo", type="primary", use_container_width=True):
            with st.spinner("Escaneando documentos..."):
                try:
                    prompt_maestro = f"""
                    Actúa como un transcriptor universitario experto. Transcribe TODO el texto de estas imágenes manteniendo la estructura original.
                    
                    REGLAS ESTRICTAS:
                    1. ESTRUCTURA: Respeta los títulos, subtítulos, listas y sangrías.
                    2. SÍMBOLOS: Usa caracteres normales para flechas (→) y grados (°). ESTÁ TOTALMENTE PROHIBIDO usar LaTeX (como $\\rightarrow$) para texto normal. Reserva el formato LaTeX EXCLUSIVAMENTE para ecuaciones matemáticas complejas.
                    3. ESQUEMAS/CUADROS: Transcribe su contenido de forma lógica y estructurada. ESTÁ PROHIBIDO dejar marcas indicando que falta una imagen o esquema.
                    4. NOTAS: Si hay post-its o notas al margen, transcríbelas agregando "[NOTA]: " al inicio.
                    5. ABREVIATURAS: Usa este diccionario provisto para reemplazar las abreviaturas: {st.session_state['dicc_abreviaturas']}.
                    6. Devuelve SOLO la transcripción directa, sin comentarios extra.
                    """

                    response = client.models.generate_content(
                        model='gemini-3.6-flash',
                        contents=[prompt_maestro] + imagenes_pil_final,
                        config=types.GenerateContentConfig(
                            temperature=0.2,
                        )
                    )

                    st.session_state['ultima_transcripcion'] = response.text

                except Exception as e:
                    st.error(f"Hubo un error al procesar las imágenes: {e}")

        # Mostrar y Exportar
        if 'ultima_transcripcion' in st.session_state:
            st.success("¡Transcripción completada!")
            st.markdown("### Resultado:")
            st.container(border=True).markdown(st.session_state['ultima_transcripcion'])

            st.divider()
            st.subheader("📥 Exportar Documento")
            
            c_down1, c_down2 = st.columns(2)
            with c_down1:
                docx_bytes = generar_docx(st.session_state['ultima_transcripcion'], imagenes_ordenadas)
                st.download_button(
                    label="Descargar .docx (Word)", 
                    data=docx_bytes, 
                    file_name="Apuntes_Digitalizados.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    use_container_width=True
                )
            with c_down2:
                st.info("💡 Para tenerlo en PDF, abrí el archivo .docx descargado y dale a 'Guardar como PDF' en tu procesador de textos.")
